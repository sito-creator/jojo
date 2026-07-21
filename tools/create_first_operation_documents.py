from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "05_初回運用サンプル" / "Word"
ZIP_PATH = ROOT / "05_初回運用サンプル" / "初回運用サンプル_一式.zip"
FONT = "Hiragino Sans"


def set_font(run, size=10.5, bold=False, color="000000") -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), FONT)


def configure(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(title)
    set_font(run, size=17, bold=True, color="1F3A5F")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("株式会社MoMo / 初回運用サンプル / 作成日：2026年7月18日")
    set_font(run, size=9, color="5B677A")


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color="2E74B5")


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run)


def shade(cell, fill="F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc: Document, rows: list[list[str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    t.style = "Table Grid"
    t.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if r_idx == 0 or c_idx == 0:
                shade(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=9.2, bold=(r_idx == 0 or c_idx == 0))
    doc.add_paragraph()


def save(title: str, filename: str, sections: list[tuple[str, list[str] | list[list[str]]]]) -> None:
    doc = Document()
    configure(doc, title)
    for name, body in sections:
        heading(doc, name)
        if body and isinstance(body[0], list):
            table(doc, body)  # type: ignore[arg-type]
        else:
            for line in body:  # type: ignore[assignment]
                if line.startswith("- "):
                    bullet(doc, line[2:])
                else:
                    paragraph(doc, line)
    doc.save(OUT_DIR / filename)


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    save(
        "初回運用開始メモ",
        "01_初回運用開始メモ.docx",
        [
            ("目的", ["規程・台帳・雛形を作成しただけで終わらせず、2026年7月から初回運用を開始した証跡を残す。支援会社レビュー後の修正に備えつつ、最低限の記録を先行して残す。"]),
            ("初回運用の対象", [
                "- 規程管理台帳：規程26本のドラフト作成、支援会社レビュー依頼、今後の承認・施行・周知予定を記録する。",
                "- 稟議台帳：10万円以上支出の代表取締役承認を前提に、上場支援会社レビュー依頼をサンプルとして記録する。",
                "- 会議体台帳：経営会議で上場準備資料の確認・レビュー依頼方針を扱う想定で記録する。",
                "- 反社チェック台帳：新規取引先・外部専門家から実査を開始する前提で、チェック項目を記録する。",
                "- 内部監査台帳：2026年8月の規程整備・周知監査を第1回監査として登録する。",
            ]),
            ("運用開始時の注意", [
                "- サンプル行は本番運用の記入例であり、正式な取引先名、契約日、証跡URL、承認日へ差し替える。",
                "- 支援会社から規程本文の修正指摘を受けた場合、規程管理台帳の版数、承認日、施行日、周知日を更新する。",
                "- 反社チェックは契約書条項だけでなく、確認方法、確認者、結果、証跡URLを残す。",
                "- 経営会議又は取締役会で報告した場合、議事録URLを会議体台帳へ記録する。",
            ]),
        ],
    )

    save(
        "第1回内部監査 監査手続書",
        "02_第1回内部監査_監査手続書.docx",
        [
            ("監査概要", [
                ["項目", "内容"],
                ["監査番号", "IA-202608-001"],
                ["監査テーマ", "規程整備・周知"],
                ["監査予定日", "2026年8月20日"],
                ["対象期間", "2026年7月18日から監査実施日まで"],
                ["被監査部門", "管理部"],
                ["監査担当", "未定。自己監査とならないよう、外部専門家又は被監査業務から独立した担当者を検討する。"],
            ]),
            ("監査目的", [
                "社内規程の作成、レビュー依頼、承認準備、周知準備、台帳管理が開始されているかを確認する。",
                "上場準備において、規程が単なる文書ではなく、運用証跡と結びついて管理されているかを確認する。",
            ]),
            ("監査手続", [
                ["No", "監査手続", "確認資料", "判定"],
                ["1", "作成済み規程26本の一覧とWordファイルが存在することを確認する。", "01_regulations_word、規程管理台帳", "適合・一部不備・不備"],
                ["2", "規程管理台帳に、規程名、版数、主管部署、承認機関、ステータスが記録されていることを確認する。", "規程管理台帳", "適合・一部不備・不備"],
                ["3", "支援会社レビュー依頼メモと未確定論点リストが作成されていることを確認する。", "review_request_memo、open_items_for_review", "適合・一部不備・不備"],
                ["4", "規程の承認・施行・周知が未了の場合、未了である旨と今後の予定が台帳又はメモに記録されていることを確認する。", "規程管理台帳、初回運用開始メモ", "適合・一部不備・不備"],
                ["5", "稟議、会議体、反社チェック、与信、内部監査の各台帳に初回サンプル又は運用開始予定が記録されていることを確認する。", "初回運用サンプル入り台帳", "適合・一部不備・不備"],
                ["6", "指摘事項がある場合、改善担当、改善期限、フォロー方法を内部監査台帳へ記録する。", "内部監査台帳", "適合・一部不備・不備"],
            ]),
            ("報告先", ["監査結果は、代表取締役、経営会議、取締役会へ報告する。取締役会設置会社化前は、代表取締役及び経営会議への報告を先行する。"]),
        ],
    )

    save(
        "第1回内部監査 チェックリスト",
        "03_第1回内部監査_チェックリスト.docx",
        [
            ("チェックリスト", [
                ["No", "確認項目", "確認結果", "指摘事項", "改善期限"],
                ["1", "規程26本がWord形式で保管されている。", "OK・NG・対象外", "", ""],
                ["2", "支援会社レビュー用パッケージが作成されている。", "OK・NG・対象外", "", ""],
                ["3", "規程管理台帳に規程名、版数、主管部署、承認機関、ステータスが入力されている。", "OK・NG・対象外", "", ""],
                ["4", "稟議台帳に10万円以上支出の記入例又は実績が登録されている。", "OK・NG・対象外", "", ""],
                ["5", "経営会議又は取締役会でレビュー依頼方針を確認する予定が記録されている。", "OK・NG・対象外", "", ""],
                ["6", "新規取引先に対する反社チェックの実施方法、確認者、証跡URL欄が設定されている。", "OK・NG・対象外", "", ""],
                ["7", "与信管理台帳に取引内容、金額、支払条件、未入金額、リスク評価欄が設定されている。", "OK・NG・対象外", "", ""],
                ["8", "内部監査台帳に第1回監査の予定が登録されている。", "OK・NG・対象外", "", ""],
                ["9", "未確定論点が、支援会社へ確認すべき内容として整理されている。", "OK・NG・対象外", "", ""],
                ["10", "指摘事項が発生した場合の改善担当、改善期限、フォロー結果欄が用意されている。", "OK・NG・対象外", "", ""],
            ]),
            ("監査人メモ", ["確認した資料、口頭確認事項、判断理由、追加確認依頼事項を記載する。", ""]),
        ],
    )

    print(f"docs={len(list(OUT_DIR.glob('*.docx')))}")


if __name__ == "__main__":
    main()

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "06_機関設計変更" / "Word"
OUT_PATH = OUT_DIR / "機関設計変更_支援会社確認メモ.docx"
FONT = "Hiragino Sans"


def set_font(run, size=11, bold=False, color="000000") -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), FONT)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("機関設計変更 支援会社確認メモ")
    set_font(run, size=18, bold=True, color="1F3A5F")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("株式会社MoMo / 作成日：2026年7月18日")
    set_font(run, size=9, color="5B677A")


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color="2E74B5")


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run)


def shade(cell, fill="F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc: Document, rows: list[list[str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    t.style = "Table Grid"
    t.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if r_idx == 0 or c_idx == 0:
                shade(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=9.5, bold=(r_idx == 0 or c_idx == 0))
    doc.add_paragraph()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    heading(doc, "1. 本メモの位置づけ")
    paragraph(doc, "本メモは、株式会社MoMoがN-1期中に取締役会設置会社・監査役設置会社へ移行するためのたたき台である。支援会社レビュー前に方針確認を行う目的で作成しており、最終的な法務・登記判断は顧問専門家及び司法書士に確認する。")

    heading(doc, "2. 現時点の前提")
    table(doc, [
        ["項目", "内容"],
        ["現在の会社形態", "取締役会非設置会社。取締役2名、監査役なし。"],
        ["定款", "最新定款あり。取締役会・監査役に関する特段の定めはないと思われるため、内容確認が必要。"],
        ["役員候補", "追加取締役は内部候補を優先。監査役は内部又は外部で検討。"],
        ["実施時期", "N-1期中を目標。"],
        ["株主総会", "株主は伊藤翔のみ。臨時株主総会はすぐ可能。"],
        ["専門家", "司法書士はいない。顧問専門家はいる。登記実務の依頼先を探す必要あり。"],
        ["支援会社確認方針", "支援会社レビュー前に本たたき台を見せる。"],
    ])

    heading(doc, "3. 変更方針案")
    bullet(doc, "取締役会設置会社へ移行するため、取締役3名体制を前提に追加取締役候補を選定する。")
    bullet(doc, "監査役設置会社へ移行するため、監査役候補を内部又は外部から選定する。上場準備の観点では、独立性、会計・法務知見、取締役会での発言力を確認する。")
    bullet(doc, "定款に、取締役会設置、監査役設置、取締役会の招集・決議、監査役に関する規定を追加する方向で検討する。")
    bullet(doc, "臨時株主総会で、定款変更、追加取締役選任、監査役選任をまとめて決議できるか確認する。")
    bullet(doc, "決議後、初回取締役会を開催し、代表取締役、会議体運用、規程承認、月次報告体制を確認する。")
    bullet(doc, "登記完了後、履歴事項全部証明書を取得し、支援会社・規程管理台帳・内部監査資料に反映する。")

    heading(doc, "4. 支援会社へ確認したい事項")
    table(doc, [
        ["No", "確認事項", "希望する回答"],
        ["1", "N-1期中に取締役会設置会社・監査役設置会社へ移行するタイミングは妥当か。", "推奨時期、優先順位、上場準備上の注意点"],
        ["2", "追加取締役は内部候補でよいか、外部取締役候補も検討すべきか。", "候補者要件、管掌領域、選定基準"],
        ["3", "監査役は内部候補でもよいか、外部候補を優先すべきか。", "独立性、常勤/非常勤、候補者要件"],
        ["4", "定款に追加すべき条項、不要な条項、上場準備上入れておきたい条項は何か。", "定款レビュー方針"],
        ["5", "株主が伊藤翔のみの場合の臨時株主総会手続はどこまで簡略化できるか。", "招集省略同意、議事録、決議書の扱い"],
        ["6", "司法書士、顧問専門家、支援会社の役割分担をどうするべきか。", "レビュー順序、依頼範囲"],
    ])

    heading(doc, "5. 直近アクション")
    table(doc, [
        ["優先", "アクション", "担当", "成果物"],
        ["1", "最新定款・登記簿・株主名簿を確認する。", "管理部", "確認メモ"],
        ["2", "追加取締役候補・監査役候補の要件を支援会社へ確認する。", "伊藤翔", "候補者要件表"],
        ["3", "司法書士候補を探し、登記に必要な手続と書類を確認する。", "管理部", "見積・必要書類リスト"],
        ["4", "定款変更案と株主総会議事録案を作成する。", "管理部/専門家", "定款案、議事録案"],
        ["5", "株主総会・初回取締役会・登記申請のスケジュールを確定する。", "伊藤翔/管理部", "確定スケジュール"],
    ])

    heading(doc, "6. 参考にした公式情報")
    paragraph(doc, "会社法及び法務省の株式会社設立手続に関する公開情報では、取締役会設置会社の場合の取締役3名以上の前提、監査役設置会社における監査役、ならびに取締役会設置会社・監査役設置会社である旨等が登記事項となることが示されている。既存会社の機関設計変更における具体的な登記書類・期限・添付書類は、司法書士に確認する。")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()

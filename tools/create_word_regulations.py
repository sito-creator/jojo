from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "01_規程ドラフト"
OUT_DIR = ROOT / "02_Word規程ドラフト"
ZIP_PATH = ROOT / "02_Word規程ドラフト.zip"

PRESET = {
    "base_font": "Arial Unicode MS",
    "jp_font": "Arial Unicode MS",
    "body_size": 11,
    "body_after": 6,
    "body_line": 1.10,
    "h1_size": 16,
    "h2_size": 13,
    "h3_size": 12,
    "h1_color": "2E74B5",
    "h3_color": "1F4D78",
    "table_header_fill": "F2F4F7",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    set_cell_margins(table)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = PRESET["base_font"]
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), PRESET["base_font"])
    r_fonts.set(qn("w:hAnsi"), PRESET["base_font"])
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), PRESET["base_font"])


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    set_east_asia_font(run, PRESET["jp_font"])
    run.bold = bold
    run.italic = italic
    return run


def apply_inline_markdown(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_run(paragraph, part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = add_run(paragraph, part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            add_run(paragraph, part)


def set_paragraph_spacing(paragraph, after_pt: int = 6, line_spacing: float = 1.10) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = PRESET["base_font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(qn("w:ascii"), PRESET["base_font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["base_font"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["jp_font"])
    normal._element.rPr.rFonts.set(qn("w:cs"), PRESET["base_font"])
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line"]

    for style_name, size, color, before, after in [
        ("Heading 1", 16, PRESET["h1_color"], 16, 8),
        ("Heading 2", 13, PRESET["h1_color"], 12, 6),
        ("Heading 3", 12, PRESET["h3_color"], 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = PRESET["base_font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["base_font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["base_font"])
        style._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["jp_font"])
        style._element.rPr.rFonts.set(qn("w:cs"), PRESET["base_font"])
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(doc: Document, label: str) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    set_east_asia_font(run, PRESET["jp_font"])
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def table_widths(column_count: int) -> list[int]:
    if column_count <= 0:
        return [9360]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = table_widths(col_count)
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            text = row[col_idx] if col_idx < len(row) else ""
            para = cell.paragraphs[0]
            set_paragraph_spacing(para, after_pt=0, line_spacing=1.10)
            apply_inline_markdown(para, text)
            for run in para.runs:
                run.font.size = Pt(9)
                if row_idx == 0:
                    run.bold = True
            if row_idx == 0:
                set_cell_shading(cell, PRESET["table_header_fill"])

    doc.add_paragraph()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and all(ch in "|:- " for ch in stripped) and "-" in stripped


def parse_table_line(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def add_body_paragraph(doc: Document, text: str) -> None:
    bullet = re.match(r"^[-*]\s+(.+)$", text)
    ordered = re.match(r"^(\d+)\.\s+(.+)$", text)
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_paragraph_spacing(p, after_pt=8, line_spacing=1.167)
        apply_inline_markdown(p, bullet.group(1))
    elif ordered:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_paragraph_spacing(p, after_pt=8, line_spacing=1.167)
        apply_inline_markdown(p, ordered.group(2))
    else:
        p = doc.add_paragraph()
        set_paragraph_spacing(p)
        apply_inline_markdown(p, text)


def convert_markdown(md_path: Path, out_path: Path) -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    first_heading = None
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush_table()
            continue
        if line.lstrip().startswith("|") and "|" in line:
            if is_table_separator(line):
                continue
            table_buffer.append(parse_table_line(line))
            continue
        flush_table()

        if line.startswith("# "):
            title = line[2:].strip()
            if first_heading is None:
                first_heading = title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(14)
                run = p.add_run(title)
                set_east_asia_font(run, PRESET["jp_font"])
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor.from_string(PRESET["h1_color"])
            else:
                p = doc.add_paragraph(style="Heading 1")
                apply_inline_markdown(p, title)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            apply_inline_markdown(p, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            apply_inline_markdown(p, line[4:].strip())
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            apply_inline_markdown(p, line[5:].strip())
        else:
            add_body_paragraph(doc, line.strip())

    flush_table()
    add_footer(doc, first_heading or md_path.stem)
    doc.core_properties.title = first_heading or md_path.stem
    doc.core_properties.author = "株式会社MoMo"
    doc.save(out_path)


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    generated: list[Path] = []
    for md_path in sorted(SRC_DIR.glob("*_ドラフト.md")):
        out_name = md_path.name.replace(".md", ".docx")
        out_path = OUT_DIR / out_name
        convert_markdown(md_path, out_path)
        generated.append(out_path)

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in generated:
            zf.write(path, path.relative_to(ROOT))

    print(f"generated={len(generated)}")
    print(f"out_dir={OUT_DIR}")
    print(f"zip={ZIP_PATH}")


if __name__ == "__main__":
    main()

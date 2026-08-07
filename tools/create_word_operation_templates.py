from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "03_運用テンプレート" / "Word雛形"
ZIP_PATH = ROOT / "03_運用テンプレート" / "Word雛形_一式.zip"

BASE_FONT = "Arial Unicode MS"
JP_FONT = "Arial Unicode MS"


def set_run_font(run, size=10.5, bold=False, color="000000") -> None:
    run.font.name = BASE_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), JP_FONT)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(title)
    set_run_font(run, size=16, bold=True, color="1F3A5F")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("株式会社MoMo　管理部")
    set_run_font(run, size=9, color="5B677A")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color="2E74B5")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths and c_idx < len(widths):
                cell.width = Inches(widths[c_idx])
            if c_idx == 0 or r_idx == 0:
                shade_cell(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=9.5, bold=(c_idx == 0 or r_idx == 0), color="000000")
    doc.add_paragraph()


def add_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, size=10.5)


def save_doc(filename: str, title: str, sections: list[tuple[str, list[list[str]] | list[str]]]) -> None:
    doc = Document()
    configure_doc(doc, title)
    for heading, body in sections:
        add_heading(doc, heading)
        if body and isinstance(body[0], list):
            add_table(doc, body)  # type: ignore[arg-type]
        else:
            add_lines(doc, body)  # type: ignore[arg-type]
    doc.save(OUT_DIR / filename)


TEMPLATES = [
    (
        "01_取締役会議事録_テンプレート.docx",
        "取締役会議事録テンプレート",
        [
            ("基本情報", [["開催日時", "YYYY年MM月DD日 HH:MM-HH:MM"], ["開催場所/方法", "本社会議室・オンライン等"], ["出席取締役", ""], ["出席監査役", ""], ["議長", "代表取締役 伊藤翔"], ["議事録作成者", "管理部"]]),
            ("決議事項", [["議案番号", "議案名", "審議内容", "決議結果", "特別利害関係の有無"], ["1", "", "", "承認・否決・継続審議", "有・無"]]),
            ("報告事項", [["報告番号", "報告事項", "報告者", "主な質疑・指摘", "次回対応"], ["1", "月次決算・予実差異", "", "", ""]]),
            ("確認欄", [["議事録確定日", ""], ["署名/電子署名", ""], ["保管場所URL", ""], ["次回取締役会予定", ""]]),
        ],
    ),
    (
        "02_経営会議議事録_テンプレート.docx",
        "経営会議議事録テンプレート",
        [
            ("基本情報", [["開催日時", ""], ["出席者", ""], ["議長", "代表取締役 伊藤翔"], ["事務局", "管理部"], ["対象月", "YYYY年MM月"]]),
            ("アジェンダ", [["番号", "議題", "担当", "決定事項", "期限"], ["1", "予実差異・KPI確認", "", "", ""], ["2", "重要案件・契約・採用", "", "", ""]]),
            ("フォロー事項", [["前回宿題", "担当", "期限", "進捗", "次アクション"], ["", "", "", "", ""]]),
            ("確認欄", [["議事録作成日", ""], ["確定日", ""], ["保管場所URL", ""], ["取締役会報告要否", "要・不要"]]),
        ],
    ),
    (
        "03_稟議書_テンプレート.docx",
        "稟議書テンプレート",
        [
            ("申請情報", [["申請日", ""], ["申請部署", ""], ["申請者", ""], ["件名", ""], ["区分", "支出・契約・採用・外注・その他"], ["金額", "円"], ["予算内外", "予算内・予算外"]]),
            ("申請内容", ["目的：", "背景：", "選定理由：", "期待効果：", "リスク・留意点："]),
            ("承認欄", [["一次確認", ""], ["法務確認", "要・不要"], ["経理確認", "要・不要"], ["最終承認者", "伊藤翔"], ["承認日", ""], ["証跡URL", ""]]),
        ],
    ),
    (
        "04_反社チェック記録票_テンプレート.docx",
        "反社会的勢力チェック記録票テンプレート",
        [
            ("対象情報", [["確認日", ""], ["取引先名", ""], ["代表者名", ""], ["所在地", ""], ["取引区分", "顧客・購買先・外注先・その他"], ["新規/既存", "新規・既存"]]),
            ("確認内容", [["確認方法", "Web検索・新聞検索・専門DB・紹介元確認等"], ["検索キーワード", ""], ["確認結果", "問題なし・要確認・取引不可・保留"], ["契約書反社条項", "有・無"], ["確認者", ""], ["レビュー者", ""]]),
            ("証跡", [["検索結果保存場所URL", ""], ["次回確認日", ""], ["備考", ""]]),
        ],
    ),
    (
        "05_与信審査票_テンプレート.docx",
        "与信審査票テンプレート",
        [
            ("取引概要", [["確認日", ""], ["取引先名", ""], ["事業区分", "AI研修・AI伴走・OEM・AI受託開発・SaaS・その他"], ["取引内容", ""], ["契約額/月額", ""], ["支払条件", ""]]),
            ("確認項目", [["会社情報確認", "済・未"], ["過去取引/入金状況", ""], ["未入金残高", ""], ["与信限度額", ""], ["リスク評価", "高・中・低"], ["対応方針", ""]]),
            ("承認欄", [["確認者", ""], ["レビュー者", ""], ["承認日", ""], ["証跡URL", ""]]),
        ],
    ),
    (
        "06_契約審査記録票_テンプレート.docx",
        "契約審査記録票テンプレート",
        [
            ("契約情報", [["取引先名", ""], ["契約名", ""], ["契約類型", "NDA・業務委託・SaaS・開発委託・その他"], ["契約期間", ""], ["契約金額", ""], ["締結予定日", ""]]),
            ("レビュー項目", [["反社条項", "有・無"], ["秘密保持", "問題なし・要修正"], ["個人情報", "対象外・要確認"], ["知的財産/成果物", "問題なし・要修正"], ["損害賠償/解除", "問題なし・要修正"], ["その他重要条項", ""]]),
            ("承認欄", [["レビュー担当", "管理部（塚原・業務委託が確認支援）"], ["最終署名者", "伊藤翔"], ["承認日", ""], ["契約書保管場所URL", ""]]),
        ],
    ),
    (
        "07_内部監査調書_テンプレート.docx",
        "内部監査調書テンプレート",
        [
            ("監査概要", [["監査対象", ""], ["監査領域", "規程・稟議・販売・購買・外注・月次決算・反社・与信等"], ["監査日", ""], ["監査担当", ""], ["被監査部門", ""], ["対象期間", ""]]),
            ("確認手続", [["確認資料", "確認結果", "指摘有無", "備考"], ["規程・台帳・証跡", "適合・一部不備・不備", "有・無", ""]]),
            ("指摘事項", [["指摘番号", "指摘内容", "重要度", "改善担当", "改善期限", "フォロー結果"], ["1", "", "高・中・低", "", "", ""]]),
            ("承認欄", [["監査担当署名", ""], ["代表取締役確認", ""], ["報告日", ""], ["証跡URL", ""]]),
        ],
    ),
    (
        "08_関連当事者調査票_テンプレート.docx",
        "関連当事者調査票テンプレート",
        [
            ("対象者情報", [["対象者名", ""], ["役職/属性", "役員・主要株主・近親者・関係会社等"], ["回答日", ""], ["対象期間", ""]]),
            ("確認事項", [["関連当事者に該当する法人・個人", "有・無"], ["MoMoとの取引", "有・無"], ["役員兼任・出資・貸付等", "有・無"], ["競業・利益相反の可能性", "有・無"], ["詳細", ""]]),
            ("会社確認欄", [["確認者", ""], ["レビュー者", ""], ["取締役会付議要否", "要・不要"], ["証跡URL", ""]]),
        ],
    ),
    (
        "09_システムアカウント申請書_テンプレート.docx",
        "システムアカウント申請書テンプレート",
        [
            ("申請情報", [["申請日", ""], ["対象者", ""], ["所属/役割", ""], ["利用サービス", "Google Workspace・Notion・マネーフォワード・SaaS管理画面等"], ["申請区分", "新規・変更・削除"], ["必要権限", ""]]),
            ("確認事項", [["個人情報アクセス", "有・無"], ["管理者権限", "有・無"], ["退職/異動時削除予定", ""], ["利用目的", ""], ["リスク・留意点", ""]]),
            ("承認欄", [["申請者", ""], ["管理責任者", ""], ["承認者", ""], ["対応完了日", ""], ["証跡URL", ""]]),
        ],
    ),
    (
        "10_クレーム報告書_テンプレート.docx",
        "クレーム報告書テンプレート",
        [
            ("受付情報", [["受付日", ""], ["受付者", ""], ["顧客/取引先", ""], ["事業区分", ""], ["受付経路", "メール・電話・フォーム・その他"], ["重大度", "高・中・低"]]),
            ("内容", ["事象概要：", "一次対応：", "原因分析：", "再発防止策：", "顧客への回答内容："]),
            ("フォロー", [["対応担当", ""], ["対応期限", ""], ["完了日", ""], ["経営会議/取締役会報告要否", "要・不要"], ["証跡URL", ""]]),
        ],
    ),
]


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for filename, title, sections in TEMPLATES:
        save_doc(filename, title, sections)

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.glob("*.docx")):
            zf.write(path, path.relative_to(ROOT / "03_運用テンプレート"))

    print(f"generated={len(TEMPLATES)}")
    print(f"out_dir={OUT_DIR}")
    print(f"zip={ZIP_PATH}")


if __name__ == "__main__":
    main()

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "04_支援会社提出パッケージ"
ZIP_PATH = ROOT / "04_支援会社提出パッケージ.zip"
PYTHON_FONT = "Arial Unicode MS"


def set_font(run, size=11, bold=False, color="000000") -> None:
    run.font.name = PYTHON_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), PYTHON_FONT)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = PYTHON_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PYTHON_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(title)
    set_font(run, size=18, bold=True, color="1F3A5F")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("株式会社MoMo / 上場準備レビュー資料 / 作成日：2026年8月5日")
    set_font(run, size=9, color="5B677A")


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color="2E74B5")


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run, size=11)


def shade_cell(cell, fill="F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if r_idx == 0:
                shade_cell(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=9.5, bold=(r_idx == 0))
    doc.add_paragraph()


def create_review_memo(path: Path) -> None:
    doc = Document()
    configure_doc(doc, "上場準備資料 レビュー依頼メモ")

    heading(doc, "1. レビュー依頼の目的")
    para(doc, "株式会社MoMoの上場準備に向け、現時点で作成した社内規程、運用台帳、各種Word雛形及び未確定事項一覧について、上場支援会社の観点から不足事項、修正事項、優先順位をご確認いただきたい。")

    heading(doc, "2. 会社前提")
    add_table(doc, [
        ["項目", "内容"],
        ["会社名", "株式会社MoMo"],
        ["決算月", "6月"],
        ["想定市場", "TOKYO PRO Marketを経由し、グロース市場を目指す"],
        ["想定時期", "現在N-1期。N-1期中に取締役会設置会社・監査役設置会社へ移行する方針"],
        ["現機関設計", "取締役2名、監査役なし。追加取締役1名及び外部非常勤監査役は候補者未定"],
        ["代表取締役", "伊藤翔"],
        ["主管部署", "管理部"],
        ["主要システム", "マネーフォワード（見積・契約・請求・経費・会計・給与）、Notion、Google Workspace、Google Sheets、GitHub等"],
    ])

    heading(doc, "3. 提出物")
    add_table(doc, [
        ["区分", "内容", "格納先"],
        ["規程Word", "上場準備に必要な規程ドラフト35本", "01_regulations_word"],
        ["運用台帳", "規程管理、稟議、会議体、反社チェック、与信、契約、外注、個人情報・システム、関連当事者、固定資産、内部監査のExcel台帳", "02_operation_ledgers"],
        ["Word雛形", "議事録、稟議書、反社チェック記録票、与信審査票、契約審査票、内部監査調書等10本", "03_operation_word_templates"],
        ["進行管理", "ロードマップ、ヒアリングシート、内部監査計画、Google Sheets確認メモ、未確定事項一覧、ひな形準拠チェックリスト", "04_project_management"],
    ])

    heading(doc, "4. 重点レビュー依頼事項")
    for item in [
        "TOKYO PRO Marketを経由してグロース市場を目指す前提で、N-1期に整備すべき規程体系に過不足がないか。",
        "取締役会設置会社化、追加取締役選任、外部非常勤監査役選任、定款変更・登記の順序と時期。",
        "10万円以上代表承認、50万円以上経営会議又は取締役会報告、100万円以上原則取締役会決議という決裁基準の妥当性。",
        "Notion、マネーフォワード、Google Driveを証跡管理の中心とする初期運用の妥当性。",
        "反社チェック、与信管理、契約審査、稟議、内部監査について、初年度に最低限運用すべき証跡とサンプル数。",
        "関連当事者取引、内部通報、個人情報、情報セキュリティ、AI入力禁止情報の初期運用として不足がないか。",
        "就業規則PDFとの不整合、賃金規程、人事評価規程、賞与、固定残業10時間分について社労士レビュー前に整理すべき論点。",
        "税理士法人cosmos及び会計士候補に確認すべき月次決算、売上計上、固定資産、棚卸資産、資金繰り表の論点。",
    ]:
        bullet(doc, item)

    heading(doc, "5. こちらで想定している次アクション")
    add_table(doc, [
        ["優先", "アクション", "目的"],
        ["1", "支援会社レビューを受ける", "規程体系、機関設計、決裁基準、証跡設計の方向性を固める"],
        ["2", "監査役候補・追加取締役候補の要件を確認する", "N-1期中の機関設計変更に備える"],
        ["3", "Notion台帳及びGoogleドキュメント様式を整備する", "稟議、反社、関連当事者、文書、クレーム等の運用開始に備える"],
        ["4", "第1回内部監査の監査手続書・チェックリストを具体化する", "N-1期下期の初回内部監査に備える"],
        ["5", "社労士、弁護士、税理士、司法書士等へ専門家レビューを依頼する", "法令・税務・登記・労務の論点を確定する"],
    ])

    doc.save(path)


def create_open_items(path: Path) -> None:
    doc = Document()
    configure_doc(doc, "未確定論点・確認事項リスト")
    heading(doc, "1. 支援会社へ確認したい未確定論点")
    add_table(doc, [
        ["No", "領域", "現状反映済み事項", "確認したい内容", "優先度"],
        ["1", "機関設計", "N-1期中に取締役会設置会社・監査役設置会社へ移行方針。追加取締役1名は未定、外部候補を検討。外部非常勤監査役候補も未定。", "追加取締役候補及び外部監査役候補の要件、選任時期、報酬、任期、責任限定契約、監査役会を当面置かない設計の妥当性", "高"],
        ["2", "会議体", "月次決算完了後すぐ経営会議を開催し、必要事項を取締役会へ上程。資料提出期限は3日前。議事録等はNotion確認記録で管理。", "取締役会・経営会議の開催頻度、議題、議事録承認証跡、決議・報告事項の粒度", "高"],
        ["3", "決裁統制", "10万円以上代表承認、50万円以上経営会議又は取締役会報告、100万円以上原則取締役会決議。事後稟議は3営業日以内。", "MoMoの規模及び上場準備水準から見た金額基準、事後稟議許容範囲、通常取引の例外設定", "高"],
        ["4", "内部監査", "初期は外部専門家活用。社内候補者なし。外部専門家未定。初回内部監査はN-1期下期想定。", "外部内部監査支援先の選定要件、独立性、初年度監査範囲、初回サンプル件数、監査役監査との役割分担", "高"],
        ["5", "反社・関連当事者", "反社は管理部がWeb検索等の簡易チェックで開始。疑義相談先はOneasia越路先生。関連当事者管理責任者は未定。既存取引なし。", "反社チェック方法、既存取引先の遡及範囲、関連当事者調査票対象、近親者・主要株主範囲、承認例外、開示検討時期", "高"],
        ["6", "労務", "末締め翌20日支払い、固定残業10時間分、通常賞与あり、人事評価は半期評価で賞与連動。就業規則PDFに支払日・労働時間・賞与算定期間の不整合あり。", "社労士レビュー前に修正すべき就業規則・賃金規程・人事評価規程の論点、賞与算定期間、固定残業の記載方法", "高"],
        ["7", "経理・資金", "経理責任者及び月次確認者は柿塚。税理士は税理士法人cosmos。資金繰り表・銀行口座一覧はN-1期中に管理部が作成。", "月次決算締め日程、月次処理項目、資金繰り表フォーマット、売上計上基準、固定資産・棚卸・税務処理のレビュー順序", "中"],
        ["8", "情報管理", "情報セキュリティ責任者は張、個人情報保護責任者は柿塚。主要サービスMFA必須。プライバシーポリシーは未公開。", "プライバシーポリシー公開時期、AI入力禁止情報の詳細、ログ管理、インシデント時の連絡手段・顧客報告判断者", "中"],
        ["9", "文書・証跡", "文書管理台帳はNotion。Google Drive、Notion、マネーフォワード権限管理者は出水。契約書は電子契約中心。", "電子帳簿保存法対応、フォルダ・ページ設計、保存期間10年中心の妥当性、紙原本発生時の統制", "中"],
        ["10", "販売・購買・外注・クレーム", "MFとNotionを中心に運用。購買相見積は50万円以上目安。外注検収は発注部門責任者。クレーム管理責任者は管理部。", "承認フロー、発注担当者、外注先選定基準、返金・値引・補償・契約解除の金額基準、研修品質問題の初動責任者", "中"],
    ])

    heading(doc, "2. 支援会社へ確認したいレビュー観点")
    for item in [
        "現時点の規程35本の粒度が、TOKYO PRO Market準備段階として過不足ないか。",
        "グロース市場を見据えた場合に、今から追加すべき規程、細則、要領、業務フロー、RCMがあるか。",
        "内部統制文書化に向け、業務記述書・フローチャート・RCMの作成順序はどこから始めるべきか。",
        "監査法人・J-Adviser・主幹事候補に提出する場合、先に整えるべき証跡は何か。",
    ]:
        bullet(doc, item)

    doc.save(path)


def copy_tree(src: Path, dst: Path) -> None:
    if dst.exists():
        shutil.rmtree(dst)
    shutil.copytree(src, dst)


def create_package() -> None:
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    PACKAGE.mkdir(parents=True, exist_ok=True)

    (PACKAGE / "00_review_request").mkdir()
    (PACKAGE / "01_regulations_word").mkdir()
    (PACKAGE / "02_operation_ledgers").mkdir()
    (PACKAGE / "03_operation_word_templates").mkdir()
    (PACKAGE / "04_project_management").mkdir()

    create_review_memo(PACKAGE / "00_review_request" / "review_request_memo.docx")
    create_open_items(PACKAGE / "00_review_request" / "open_items_for_review.docx")

    copy_tree(ROOT / "02_Word規程ドラフト_ASCIIファイル名", PACKAGE / "01_regulations_word")
    shutil.copy2(ROOT / "03_運用テンプレート" / "Excel台帳" / "上場準備_運用台帳セット.xlsx", PACKAGE / "02_operation_ledgers" / "ipo_operation_ledgers.xlsx")
    copy_tree(ROOT / "03_運用テンプレート" / "Word雛形", PACKAGE / "03_operation_word_templates")

    for src, name in [
        (ROOT / "00_進行管理" / "規程作成ロードマップ.md", "roadmap.md"),
        (ROOT / "00_進行管理" / "初回ヒアリングシート.md", "initial_hearing_sheet.md"),
        (ROOT / "00_進行管理" / "初年度内部監査計画_たたき台.md", "initial_internal_audit_plan.md"),
        (ROOT / "00_進行管理" / "Googleスプレッドシート確認メモ.md", "google_sheets_review_memo.md"),
        (ROOT / "00_進行管理" / "未確定事項一覧.md", "open_items_current.md"),
        (ROOT / "00_進行管理" / "未確定事項一覧.docx", "open_items_current.docx"),
        (ROOT / "00_進行管理" / "規程ひな形準拠チェックリスト.md", "template_conformance_checklist.md"),
        (ROOT / "00_進行管理" / "規程ひな形読み取り_ヒアリング整理.md", "template_reading_and_hearing_notes.md"),
    ]:
        if src.exists():
            shutil.copy2(src, PACKAGE / "04_project_management" / name)

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(PACKAGE.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(PACKAGE.parent))

    print(f"package={PACKAGE}")
    print(f"zip={ZIP_PATH}")
    print(f"files={sum(1 for p in PACKAGE.rglob('*') if p.is_file())}")


if __name__ == "__main__":
    create_package()
